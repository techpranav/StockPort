from docx import Document
from datetime import datetime
import markdown
from bs4 import BeautifulSoup
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re

def generate_word_report(stock_symbol, data, gpt_summary, filename):
    """Generate a well-formatted Word document for the stock analysis."""
    doc = Document()

    # ===== 1. Title =====
    title = doc.add_heading(f"{stock_symbol} - Stock Analysis Report", level=1)

    info = data.get("info", {})
    news = data.get("news", [])

    # ===== 2. Company Overview =====
    doc.add_heading("Company Overview", level=2)
    doc.add_paragraph().add_run(f"Name: ").add_text( f"{info.get('longName', stock_symbol)}")
    doc.add_paragraph().add_run(f"Business Summary: ").add_text( f"{info.get('longBusinessSummary', 'N/A')}")
    doc.add_paragraph().add_run(f"Sector: ").add_text( f"{info.get('sector', 'N/A')}")
    doc.add_paragraph().add_run(f"Industry: ").add_text( f"{info.get('industry', 'N/A')}")
    doc.add_paragraph().add_run(f"Market Cap: ").add_text( f"{info.get('marketCap', 'N/A')}")
    doc.add_paragraph().add_run(f"PE Ratio: ").add_text( f"{info.get('trailingPE', 'N/A')}")
    doc.add_paragraph().add_run(f"Dividend Yield: ").add_text( f"{info.get('dividendYield', 'N/A')}")
    doc.add_paragraph().add_run(f"Website : ").add_text( f"{info.get('website', 'N/A')}")


    # ===== 3. GPT Analysis Summary =====
    doc.add_heading("Investment Summary", level=2)

    markdown_to_docx(gpt_summary,doc)

    # ===== 4. Latest News =====
    doc.add_heading("Recent News", level=2)
    if not news:
        doc.add_paragraph("No recent news available.")
    else:
        for newsElem in news:
            article = newsElem.get('content')
            pub_time = article.get("pubDate", "Unknown Date")

            # Convert from string to datetime object
            if pub_time != "Unknown Date":
                pub_time = datetime.strptime(pub_time, "%Y-%m-%dT%H:%M:%SZ").strftime("%Y-%m-%d %H:%M:%S")

            # Fix: Apply "List Bullet" correctly
            p = doc.add_paragraph(f" {article.get('title', 'No Title')}")
            p.style = "List Bullet"

            doc.add_paragraph(f"  {article.get('summary', 'Unknown')}")
            doc.add_paragraph(f"  Source: {article.get('provider', {}).get('displayName', 'Unknown')}")
            doc.add_paragraph(f"  Published: {pub_time}")
            doc.add_paragraph(f"  URL: {article.get('canonicalUrl', {}).get('url', 'Unknown')}\n")

    # ===== 5. Save Report =====
    doc.save(filename)
    return filename


def add_hyperlink(paragraph, url, text):
    """Adds a clickable hyperlink to the Word document."""
    run = paragraph.add_run(text)
    run.font.color.rgb = (0, 0, 255)  # Blue color
    run.underline = True
    part = paragraph._p
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), url)
    part.append(hyperlink)


def process_paragraph(doc, elem):
    """Processes a paragraph with mixed formatting (bold, italic, inline code)."""
    p = doc.add_paragraph()
    for part in elem.contents:
        if part.name == "strong":  # Bold
            run = p.add_run(part.text)
            run.bold = True
        elif part.name == "em":  # Italic
            run = p.add_run(part.text)
            run.italic = True
        elif part.name == "code":  # Inline Code
            run = p.add_run(part.text)
            run.font.name = "Courier New"
        else:
            run = p.add_run(part)  # Normal text


def markdown_to_docx(md_text, doc):
    """Converts Markdown to a properly styled Word document."""

    # Convert Markdown to HTML
    html_text = markdown.markdown(md_text)

    # Parse HTML
    soup = BeautifulSoup(html_text, "html.parser")



    # Process parsed elements
    for elem in soup.children:
        if elem.name in ["h1", "h2", "h3", "h4", "h5", "h6"]:
            level = int(elem.name[1])  # Extract heading level (h1 → 1, h2 → 2)
            doc.add_paragraph(elem.text, style=f"Heading {level}")

        elif elem.name == "p":
            process_paragraph(doc, elem)

        elif elem.name == "pre":  # Code Block
            code_text = elem.text.strip()
            p = doc.add_paragraph()
            run = p.add_run(code_text)
            run.font.name = "Courier New"

        elif elem.name == "ul":  # Unordered List
            for li in elem.find_all("li"):
                doc.add_paragraph(li.text, style="ListBullet")

        elif elem.name == "ol":  # Ordered List
            for li in elem.find_all("li"):
                text = re.sub(r'^\d+\.?\s*', '', li.text)
                doc.add_paragraph(text, style="ListNumber")


        elif elem.name == "blockquote":  # Blockquote
            p = doc.add_paragraph()
            run = p.add_run(elem.text)
            p.paragraph_format.left_indent = Pt(20)  # Indent blockquotes
            run.italic = True

        elif elem.name == "a":  # Links
            p = doc.add_paragraph()
            add_hyperlink(p, elem["href"], elem.text)

    return doc

