from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import markdown
from bs4 import BeautifulSoup
import re
from datetime import datetime
from typing import Dict, Any, List
from pathlib import Path
import pandas as pd
from core.config import OUTPUT_DIR, ENABLE_AI_FEATURES
from constants.StringConstants import (
    FINANCIAL_STATEMENT_FILTER_KEYS,
    FINANCIAL_SHEET_NAMES,
    FINANCIAL_STATEMENT_TYPES,
    FINANCIAL_METRICS_KEYS,
    BALANCE_SHEET_EXPORT_KEYS,
    INCOME_STATEMENT_EXPORT_KEYS,
    CASHFLOW_EXPORT_KEYS,
    # File and directory constants
    REPORTS_DIR,
    EXCEL_EXTENSION,
    WORD_EXTENSION,
    REPORT_FILE_PATTERN,
    ANALYSIS_REPORT_FILE_PATTERN,
    # Sheet names
    SHEET_COMPANY_INFO,
    SHEET_KEY_METRICS,
    SHEET_TECHNICAL_ANALYSIS,
    SHEET_FUNDAMENTAL_ANALYSIS,
    SHEET_PORTFOLIO_ANALYSIS,
    # Column names
    COLUMN_METRIC,
    COLUMN_VALUE,
    COLUMN_ATTRIBUTE,
    COLUMN_INDICATOR,
    # Formatting
    NOT_AVAILABLE
)
from constants.Messages import (
    # Word document headers
    DOC_HEADER_COMPANY_INFO,
    DOC_HEADER_INCOME_STATEMENT,
    DOC_HEADER_BALANCE_SHEET,
    DOC_HEADER_CASH_FLOW,
    DOC_HEADER_AI_SUMMARY,
    # Messages
    MSG_NO_DATA_AVAILABLE
)
import os

class ReportService:
    """Service for generating stock analysis reports."""
    
    def __init__(self):
        """Initialize the report service."""
        self.reports_dir = REPORTS_DIR
        if not os.path.exists(self.reports_dir):
            os.makedirs(self.reports_dir)
    
    def generate_excel_report(self, symbol: str, data: Dict[str, Any]) -> str:
        """Generate Excel report for stock analysis."""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        excel_path = os.path.join(self.reports_dir, f"{REPORT_FILE_PATTERN.format(symbol=symbol, timestamp=timestamp)}{EXCEL_EXTENSION}")
        with pd.ExcelWriter(excel_path) as writer:
            # Write company information
            if 'info' in data:
                self._write_company_info(writer, data['info'])
            
            # Write key metrics
            if 'metrics' in data:
                self._write_metrics(writer, data['metrics'])
            
            # Write technical analysis
            if 'technical_analysis' in data:
                self._write_technical_analysis(writer, data['technical_analysis'])
            
            # Write fundamental analysis
            if 'fundamental_analysis' in data:
                self._write_fundamental_analysis(writer, data['fundamental_analysis'])
            
            # Write portfolio analysis
            if 'portfolio_analysis' in data:
                self._write_portfolio_analysis(writer, data['portfolio_analysis'])
            
            # Write financial statements
            if 'financials' in data:
                self._write_financial_statements(writer, data['financials'])
        
        return excel_path

    def _write_company_info(self, writer: pd.ExcelWriter, info: Dict[str, Any]) -> None:
        """Write company information to Excel."""
        # Flatten nested dictionaries
        flat_info = {}
        for key, value in info.items():
            if isinstance(value, dict):
                for sub_key, sub_value in value.items():
                    flat_info[f"{key}_{sub_key}"] = sub_value
            else:
                flat_info[key] = value
        
        # Convert to DataFrame
        df = pd.DataFrame(list(flat_info.items()), columns=[COLUMN_METRIC, COLUMN_VALUE])
        
        # Clean and format for Excel
        cleaned_df = self._clean_dataframe_for_excel(df, SHEET_COMPANY_INFO)
        cleaned_df.to_excel(writer, sheet_name=SHEET_COMPANY_INFO, index=False)

    def _write_metrics(self, writer: pd.ExcelWriter, metrics: Dict[str, Any]) -> None:
        """Write key metrics to Excel."""
        # Convert metrics to DataFrame
        df = pd.DataFrame(list(metrics.items()), columns=[COLUMN_METRIC, COLUMN_VALUE])
        
        # Format numeric values
        for idx, row in df.iterrows():
            if isinstance(row[COLUMN_VALUE], (int, float)) and not pd.isna(row[COLUMN_VALUE]):
                if abs(row[COLUMN_VALUE]) > 1000000:
                    df.at[idx, COLUMN_VALUE] = f"{row[COLUMN_VALUE]:,.0f}"
                elif abs(row[COLUMN_VALUE]) > 1000:
                    df.at[idx, COLUMN_VALUE] = f"{row[COLUMN_VALUE]:,.2f}"
        
        # Clean and export
        cleaned_df = self._clean_dataframe_for_excel(df, SHEET_KEY_METRICS)
        cleaned_df.to_excel(writer, sheet_name=SHEET_KEY_METRICS, index=False)
    
    def _write_technical_analysis(self, writer: pd.ExcelWriter, analysis: Dict[str, Any]) -> None:
        """Write technical analysis to Excel."""
        df = pd.DataFrame(list(analysis.items()), columns=[COLUMN_INDICATOR, COLUMN_VALUE])
        
        # Clean and format for Excel
        cleaned_df = self._clean_dataframe_for_excel(df, SHEET_TECHNICAL_ANALYSIS)
        cleaned_df.to_excel(writer, sheet_name=SHEET_TECHNICAL_ANALYSIS, index=False)
    
    def _write_fundamental_analysis(self, writer: pd.ExcelWriter, analysis: Dict[str, Any]) -> None:
        """Write fundamental analysis to Excel."""
        df = pd.DataFrame(list(analysis.items()), columns=[COLUMN_METRIC, COLUMN_VALUE])
        
        # Clean and format for Excel
        cleaned_df = self._clean_dataframe_for_excel(df, SHEET_FUNDAMENTAL_ANALYSIS)
        cleaned_df.to_excel(writer, sheet_name=SHEET_FUNDAMENTAL_ANALYSIS, index=False)
    
    def _write_portfolio_analysis(self, writer: pd.ExcelWriter, analysis: Dict[str, Any]) -> None:
        """Write portfolio analysis to Excel."""
        df = pd.DataFrame(list(analysis.items()), columns=[COLUMN_METRIC, COLUMN_VALUE])
        
        # Clean and format for Excel
        cleaned_df = self._clean_dataframe_for_excel(df, SHEET_PORTFOLIO_ANALYSIS)
        cleaned_df.to_excel(writer, sheet_name=SHEET_PORTFOLIO_ANALYSIS, index=False)
    
    def _clean_dataframe_for_excel(self, df: pd.DataFrame, sheet_name: str = "") -> pd.DataFrame:
        """Clean DataFrame for Excel export by handling NaN values and formatting."""
        if df is None or df.empty:
            return pd.DataFrame()
        
        try:
            # Create a copy to avoid modifying original
            cleaned_df = df.copy()
            
            # Handle different data types
            for col in cleaned_df.columns:
                if cleaned_df[col].dtype in ['float64', 'int64']:
                    # For numeric columns, replace NaN with 0 and format large numbers
                    cleaned_df[col] = cleaned_df[col].fillna(0)
                    
                    # Format large numbers for better readability
                    if cleaned_df[col].abs().max() > 1000000:
                        # For millions, keep as numbers but round to 2 decimal places
                        cleaned_df[col] = cleaned_df[col].round(2)
                    elif cleaned_df[col].abs().max() > 1000:
                        # For thousands, round to 2 decimal places
                        cleaned_df[col] = cleaned_df[col].round(2)
                    else:
                        # For smaller numbers, round to 4 decimal places
                        cleaned_df[col] = cleaned_df[col].round(4)
                        
                elif cleaned_df[col].dtype == 'object':
                    # For text/object columns, replace NaN with empty string
                    cleaned_df[col] = cleaned_df[col].fillna('')
                    
                    # Clean up any remaining NaN-like strings
                    cleaned_df[col] = cleaned_df[col].replace(['nan', 'NaN', 'None', 'null'], '')
                    
                else:
                    # For other data types, replace NaN with empty string
                    cleaned_df[col] = cleaned_df[col].fillna('')
            
            # Clean up index if it contains timestamps or other problematic data
            if hasattr(cleaned_df.index, 'dtype'):
                if 'datetime' in str(cleaned_df.index.dtype) or 'timestamp' in str(cleaned_df.index.dtype):
                    # Convert datetime index to string
                    cleaned_df.index = cleaned_df.index.astype(str)
            
            return cleaned_df
            
        except Exception as e:
            print(f"Error cleaning DataFrame for {sheet_name}: {str(e)}")
            # Return a simplified version if cleaning fails
            try:
                simple_df = df.fillna('')
                return simple_df
            except:
                return df  # Return original as last resort

    def _write_financial_statements(self, writer: pd.ExcelWriter, financials: Dict[str, Any]) -> None:
        """Write financial statements to Excel with proper formatting."""
        try:
            # Write yearly statements
            for statement_type in ['income_statement', 'balance_sheet', 'cashflow']:
                key = FINANCIAL_STATEMENT_FILTER_KEYS['yearly'][statement_type]
                sheet_name = FINANCIAL_SHEET_NAMES['yearly'][statement_type]
                
                if key in financials and financials[key] is not None:
                    df = financials[key]
                    if isinstance(df, pd.DataFrame) and not df.empty:
                        cleaned_df = self._clean_dataframe_for_excel(df, sheet_name)
                        cleaned_df.to_excel(writer, sheet_name=sheet_name, na_rep='', float_format='%.2f')
                    else:
                        # Create empty DataFrame with message
                        empty_df = pd.DataFrame({'Message': ['No data available']})
                        empty_df.to_excel(writer, sheet_name=sheet_name, index=False)
            
            # Write quarterly statements
            for statement_type in ['income_statement', 'balance_sheet', 'cashflow']:
                key = FINANCIAL_STATEMENT_FILTER_KEYS['quarterly'][statement_type]
                sheet_name = FINANCIAL_SHEET_NAMES['quarterly'][statement_type]
                
                if key in financials and financials[key] is not None:
                    df = financials[key]
                    if isinstance(df, pd.DataFrame) and not df.empty:
                        cleaned_df = self._clean_dataframe_for_excel(df, sheet_name)
                        cleaned_df.to_excel(writer, sheet_name=sheet_name, na_rep='', float_format='%.2f')
                    else:
                        # Create empty DataFrame with message
                        empty_df = pd.DataFrame({'Message': ['No data available']})
                        empty_df.to_excel(writer, sheet_name=sheet_name, index=False)
                        
        except Exception as e:
            print(f"Error writing financial statements: {str(e)}")
            # Create a simple error sheet
            error_df = pd.DataFrame({'Error': [f'Error writing financial data: {str(e)}']})
            error_df.to_excel(writer, sheet_name='Financial_Error', index=False)


    def generate_word_report(self, symbol: str, data: Dict[str, Any]) -> Path:
        """Generate a Word report for the stock analysis."""
        doc = Document()
        
        # Add title
        doc.add_heading(f"Stock Analysis Report: {symbol}", 0)
        
        # Add company info
        if "info" in data:
            doc.add_heading(DOC_HEADER_COMPANY_INFO, level=1)
            for key, value in data["info"].items():
                doc.add_paragraph(f"{key}: {value}")
        
        # Add financial statements
        if "income_statement" in data and not data["income_statement"].empty:
            doc.add_heading(DOC_HEADER_INCOME_STATEMENT, level=1)
            self._add_dataframe_to_doc(doc, data["income_statement"])
        
        if "balance_sheet" in data and not data["balance_sheet"].empty:
            doc.add_heading(DOC_HEADER_BALANCE_SHEET, level=1)
            self._add_dataframe_to_doc(doc, data["balance_sheet"])
        
        if "cashflow" in data and not data["cashflow"].empty:
            doc.add_heading(DOC_HEADER_CASH_FLOW, level=1)
            self._add_dataframe_to_doc(doc, data["cashflow"])
        
        # Add AI summary if available
        if ENABLE_AI_FEATURES and "summary" in data and data["summary"]:
            doc.add_heading(DOC_HEADER_AI_SUMMARY, level=1)
            doc.add_paragraph(data["summary"])
        
        # Save the document
        output_path = os.path.join(self.reports_dir, f"{ANALYSIS_REPORT_FILE_PATTERN.format(symbol=symbol)}{WORD_EXTENSION}")

        doc.save(str(output_path))
        
        return output_path
    
    def _add_dataframe_to_doc(self, doc: Document, df: pd.DataFrame) -> None:
        """Add a DataFrame to a Word document."""
        if df.empty:
            doc.add_paragraph(MSG_NO_DATA_AVAILABLE)
            return
        
        # Convert DataFrame to table
        table = doc.add_table(rows=1, cols=len(df.columns) + 1)
        table.style = 'Table Grid'
        
        # Add headers
        header_cells = table.rows[0].cells
        header_cells[0].text = "Metric"
        for i, col in enumerate(df.columns):
            header_cells[i + 1].text = str(col)
        
        # Add data
        for idx, row in df.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = str(idx)
            for i, value in enumerate(row):
                row_cells[i + 1].text = str(value)

    def _add_company_overview(self, doc: Document, data: Dict[str, Any]) -> None:
        """Add company overview section to the document."""
        doc.add_heading("Company Overview", level=2)
        info = data.get("info", {})
        
        overview_fields = [
            ("Name", "longName"),
            ("Business Summary", "longBusinessSummary"),
            ("Sector", "sector"),
            ("Industry", "industry"),
            ("Market Cap", "marketCap"),
            ("PE Ratio", "trailingPE"),
            ("Dividend Yield", "dividendYield"),
            ("Website", "website")
        ]

        for label, key in overview_fields:
            value = info.get(key, 'N/A')
            doc.add_paragraph().add_run(f"{label}: ").add_text(str(value))

    def _add_news_section(self, doc: Document, data: Dict[str, Any]) -> None:
        """Add news section to the document."""
        doc.add_heading("Recent News", level=2)
        news = data.get("news", [])

        if not news:
            doc.add_paragraph("No recent news available.")
            return

        for news_elem in news:
            article = news_elem.get('content', {})
            pub_time = article.get("pubDate", "Unknown Date")

            if pub_time != "Unknown Date":
                pub_time = datetime.strptime(pub_time, "%Y-%m-%dT%H:%M:%SZ").strftime("%Y-%m-%d %H:%M:%S")

            p = doc.add_paragraph(f" {article.get('title', 'No Title')}")
            p.style = "List Bullet"

            doc.add_paragraph(f"  {article.get('summary', 'Unknown')}")
            doc.add_paragraph(f"  Source: {article.get('provider', {}).get('displayName', 'Unknown')}")
            doc.add_paragraph(f"  Published: {pub_time}")
            doc.add_paragraph(f"  URL: {article.get('canonicalUrl', {}).get('url', 'Unknown')}\n")

    def _markdown_to_docx(self, md_text: str, doc: Document) -> None:
        """Convert Markdown to Word document format."""
        html_text = markdown.markdown(md_text)
        soup = BeautifulSoup(html_text, "html.parser")

        for elem in soup.children:
            if elem.name in ["h1", "h2", "h3", "h4", "h5", "h6"]:
                level = int(elem.name[1])
                doc.add_paragraph(elem.text, style=f"Heading {level}")
            elif elem.name == "p":
                self._process_paragraph(doc, elem)
            elif elem.name == "pre":
                p = doc.add_paragraph()
                run = p.add_run(elem.text.strip())
                run.font.name = "Courier New"
            elif elem.name == "ul":
                for li in elem.find_all("li"):
                    doc.add_paragraph(li.text, style="ListBullet")
            elif elem.name == "ol":
                for li in elem.find_all("li"):
                    text = re.sub(r'^\d+\.?\s*', '', li.text)
                    doc.add_paragraph(text, style="ListNumber")
            elif elem.name == "blockquote":
                p = doc.add_paragraph()
                run = p.add_run(elem.text)
                p.paragraph_format.left_indent = Pt(20)
                run.italic = True
            elif elem.name == "a":
                p = doc.add_paragraph()
                self._add_hyperlink(p, elem["href"], elem.text)

    def _process_paragraph(self, doc: Document, elem: BeautifulSoup) -> None:
        """Process a paragraph with mixed formatting."""
        p = doc.add_paragraph()
        for part in elem.contents:
            if part.name == "strong":
                run = p.add_run(part.text)
                run.bold = True
            elif part.name == "em":
                run = p.add_run(part.text)
                run.italic = True
            elif part.name == "code":
                run = p.add_run(part.text)
                run.font.name = "Courier New"
            else:
                run = p.add_run(part)

    def _add_hyperlink(self, paragraph: Document, url: str, text: str) -> None:
        """Add a clickable hyperlink to the document."""
        run = paragraph.add_run(text)
        run.font.color.rgb = (0, 0, 255)
        run.underline = True
        part = paragraph._p
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), url)
        part.append(hyperlink) 